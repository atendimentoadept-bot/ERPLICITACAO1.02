import streamlit as st
import pandas as pd
from datetime import datetime
import os
import os                               # Para manipular caminhos de arquivos
from docx import Document               # Biblioteca principal para ler/escrever .docx [cite: 1]
from docx.shared import Pt              # Para ajustar tamanho de fontes, se necessário
from num2words import num2words        # Para converter números em extenso
import io

try:
    dados_produtos=pd.read_csv("Base de Dados.csv",sep=";")
    dados_pessoas = pd.read_csv("Base_Pessoas.csv", sep=";")
    print("Base de Dados OK")
except FileNotFoundError:
    print("Nao foi posssivel interagir com a Base de Dados")

st.sidebar.title("Menu")
pagina=st.sidebar.radio("Ir para:", ["Criar Pedido","Consultar Pedido","Consultar Produto","Cadastrar Produto","Cadastrar Pessoa","Consultar Pessoa","Formalizacao"])

from datetime import datetime # Certifique-se de que essa linha está no topo do seu arquivo!

if pagina == "Cadastrar Produto":
    st.title("Cadastro de Produtos")

    # Criamos abas para separar o cadastro Manual da Importação
    aba1, aba2 = st.tabs(["Cadastro Manual", "Importação em Massa (CSV)"])

    # --- ABA 1: CADASTRO MANUAL ---
    with aba1:
        st.info("O Valor Líquido será calculado: Custo + Impostos + Lucro")

        with st.form("form_cadastro"):
            st.subheader("1. Identificação")
            col_id_1, col_id_2 = st.columns(2)
            
            with col_id_1:
                id_sku = st.text_input("SKU / Código Interno (Obrigatório)")
                descricao = st.text_input("Descrição do Produto (Obrigatório)")
                marca = st.text_input("Marca / Fabricante")
                
            with col_id_2:
                categoria = st.selectbox("Categoria", ["Geral", "Eletrônicos", "Vestuário", "Ferramentas", "Outros"])
                fornecedor = st.selectbox("Fornecedor", ["Samsung", "Apple", "LG", "Motorola", "Outros"])
                
                c_est1, c_est2 = st.columns(2)
                estoque_atual = c_est1.number_input("Estoque Atual", min_value=0, step=1)
                estoque_minimo = c_est2.number_input("Estoque Mínimo", min_value=1, value=5)

            st.divider() 

            st.subheader("2. Custos e Precificação")
            col_fin_1, col_fin_2, col_fin_3 = st.columns(3)
            with col_fin_1:
                preco_custo = st.number_input("Preço de Custo (R$)", min_value=0.00, step=0.01)
                lucro = st.number_input("Margem de Lucro (R$)", min_value=0.00, step=0.01)
            
            with col_fin_2:
                icms = st.number_input("ICMS (R$)", min_value=0.0, step=0.01)
                ipi = st.number_input("IPI (R$)", min_value=0.0, step=0.01)
            
            with col_fin_3:
                valor_st = st.number_input("ST (R$)", min_value=0.0, step=0.01)
                ncm = st.text_input("NCM")

            botao_salvar = st.form_submit_button("Salvar Produto")

        # --- LÓGICA DE SALVAR (MANUAL) ---
        if botao_salvar:
            # 1. Validação de Campos Vazios
            erros = []
            if not id_sku: erros.append("O SKU é obrigatório.")
            if not descricao: erros.append("A Descrição é obrigatória.")
            
            # 2. TRAVA DE SKU DUPLICADO (A novidade aqui!)
            # Convertemos para string para garantir que '123' seja igual a '123'
            lista_skus = dados_produtos["id_sku"].astype(str).tolist()
            if str(id_sku) in lista_skus:
                erros.append(f"ERRO CRÍTICO: O SKU '{id_sku}' já existe no sistema!")

            if len(erros) > 0:
                for erro in erros:
                    st.error(erro)
            else:
                # Cálculo Automático
                valor_liquido = preco_custo + icms + ipi + valor_st + lucro
                
                # Criação da Linha
                nova_linha = pd.DataFrame({
                    "id_sku": [id_sku],
                    "descricao": [descricao],
                    "categoria": [categoria],
                    "marca": [marca],
                    "fornecedor": [fornecedor],
                    "ncm": [ncm],
                    "preco_custo": [preco_custo],
                    "lucro": [lucro],
                    "icms": [icms], "ipi": [ipi], "st": [valor_st],
                    "valor_liquido": [valor_liquido],
                    "estoque_atual": [estoque_atual],
                    "estoque_minimo": [estoque_minimo],
                    "ativo": [True],
                    "data_cadastro": [datetime.now().strftime("%d/%m/%Y")]
                    # Adicione aqui as outras colunas zeradas se necessário (peso, altura, etc)
                })

                # Salvar
                dados_produtos = pd.concat([dados_produtos, nova_linha], ignore_index=True)
                dados_produtos.to_csv("Base de Dados.csv", sep=";", index=False)
                st.success(f"✅ Produto {id_sku} cadastrado com sucesso!")

    # --- ABA 2: IMPORTAÇÃO EM MASSA (CSV) ---
    with aba2:
        st.header("Importar Produtos via CSV")
        
        # --- SEÇÃO DE INSTRUÇÕES (Expansível) ---
        with st.expander("📖 Leia as instruções antes de importar", expanded=False):
            st.markdown("""
            Para que a importação funcione corretamente, seu arquivo deve seguir estas regras:
            1. **Separador:** O arquivo deve ser salvo no formato **CSV (Separado por ponto e vírgula `;`)**.
            2. **Cabeçalhos:** A primeira linha deve conter exatamente os nomes das colunas (id_sku, descricao, etc).
            3. **Decimais:** Em valores de dinheiro, utilize o **ponto (.)** como separador decimal (Ex: 150.50 e não 150,50).
            4. **SKUs Únicos:** Se um SKU do arquivo já existir no banco de dados, essa linha será ignorada.
            """)
            
            # Tabela de exemplo com 5 linhas
            st.write("### Exemplo de preenchimento:")
            exemplo_dados = {
                "id_sku": ["CEL-S23-01", "CAM-IPH-15", "FONE-BT-LG", "NOTE-DELL-G15", "MOU-LOGI-M280"],
                "descricao": ["Smartphone S23", "iPhone 15 Pro", "Fone Bluetooth", "Notebook Gamer", "Mouse Sem Fio"],
                "fornecedor": ["Samsung", "Apple", "LG", "Dell", "Logitech"],
                "preco_custo": [4500.00, 7200.00, 150.00, 5800.00, 89.90],
                "lucro": [500.00, 800.00, 50.00, 600.00, 30.00],
                "valor_liquido": [5000.00, 8000.00, 200.00, 6400.00, 119.90]
            }
            st.table(exemplo_dados)
        
        st.divider()
        
        # --- ÁREA DE UPLOAD ---
        arquivo_upload = st.file_uploader("Arraste seu arquivo CSV aqui", type=["csv"])
        
        if arquivo_upload is not None:
            try:
                # Lendo o arquivo subido
                df_novo = pd.read_csv(arquivo_upload, sep=";")
                
                st.write("🔍 **Pré-visualização dos dados detectados:**")
                st.dataframe(df_novo.head()) # Mostra as primeiras 5 linhas para o usuário conferir
                
                if st.button("Confirmar Importação"):
                    # Verificação de SKUs
                    skus_existentes = dados_produtos["id_sku"].astype(str).tolist()
                    
                    # Filtra apenas o que é novo
                    df_novo_filtrado = df_novo[~df_novo["id_sku"].astype(str).isin(skus_existentes)]
                    
                    qtd_total = len(df_novo)
                    qtd_novos = len(df_novo_filtrado)
                    qtd_ignorados = qtd_total - qtd_novos
                    
                    if qtd_novos > 0:
                        # Garante que a data de cadastro seja preenchida
                        if "data_cadastro" not in df_novo_filtrado.columns:
                            df_novo_filtrado["data_cadastro"] = datetime.now().strftime("%d/%m/%Y")
                        
                        # Concatena e salva
                        dados_atualizados = pd.concat([dados_produtos, df_novo_filtrado], ignore_index=True)
                        dados_atualizados.to_csv("Base de Dados.csv", sep=";", index=False)
                        
                        st.success(f"✅ Sucesso! {qtd_novos} novos produtos adicionados.")
                        if qtd_ignorados > 0:
                            st.warning(f"⚠️ {qtd_ignorados} produtos foram ignorados porque o SKU já existia.")
                    else:
                        st.error("❌ Operação cancelada: Todos os produtos deste arquivo já existem no banco de dados.")
                        
            except Exception as e:
                st.error(f"Erro ao processar o arquivo. Verifique se o separador é ponto e vírgula (;). Detalhe: {e}")
#Tela de Consulta de Produtos
elif pagina == "Consultar Produto":
    st.title("Consulta de Produtos")

    # Criamos duas colunas: uma estreita para os filtros e uma larga para o resultado
    col_filtros, col_resultado = st.columns([1, 3])

    with col_filtros:
        st.subheader("Filtros de Busca")
        # Busca por SKU (Texto exato ou parcial)
        filtro_sku = st.text_input("Código SKU")
        
        # Busca por Descrição (Palavra-chave)
        filtro_desc = st.text_input("Descrição do Produto")
        
        st.divider()
        st.caption("Dica: A busca por descrição encontra palavras parciais (ex: 'azul' encontra 'Camiseta Azul').")

    with col_resultado:
        # Criamos uma cópia dos dados para não alterar o original durante o filtro
        df_filtrado = dados_produtos.copy()

        # Lógica de Filtro em Tempo Real
        if filtro_sku:
            # Filtra por SKU (transformando tudo em string para evitar erro)
            df_filtrado = df_filtrado[df_filtrado["id_sku"].astype(str).str.contains(filtro_sku, case=False, na=False)]
        
        if filtro_desc:
            # Filtra por Descrição (case=False ignora maiúsculas/minúsculas)
            df_filtrado = df_filtrado[df_filtrado["descricao"].astype(str).str.contains(filtro_desc, case=False, na=False)]

        # Selecionamos apenas as colunas solicitadas
        # Nota: Use os nomes exatos das colunas do seu CSV aqui
        colunas_exibicao = [
            "id_sku", 
            "descricao", 
            "fornecedor", 
            "preco_custo", 
            "lucro", 
            "valor_liquido"
        ]
        
        # Verificar se as colunas existem antes de exibir (para evitar erro de arquivo vazio)
        try:
            exibicao = df_filtrado[colunas_exibicao]
            
            # Renomear apenas para ficar bonito na tabela do usuário
            exibicao.columns = ["SKU", "DESCRIÇÃO", "FORNECEDOR", "CUSTO (R$)", "LUCRO (R$)", "VALOR LÍQUIDO (R$)"]
            
            st.subheader(f"Resultados ({len(exibicao)} encontrados)")
            
            if len(exibicao) > 0:
                st.dataframe(
                    exibicao, 
                    use_container_width=True, 
                    hide_index=True # Esconde aquela coluna de números 0, 1, 2...
                )
            else:
                st.warning("Nenhum produto encontrado com esses termos.")
                
        except KeyError as e:
            st.error(f"Erro: Alguma coluna não foi encontrada no CSV: {e}")
# --- 5. TELA DE CADASTRO DE PESSOAS ---
elif pagina == "Cadastrar Pessoa":
    st.title("Cadastro de Clientes e Fornecedores")

    # Tenta carregar a base de pessoas, se não existir, cria uma vazia
    try:
        dados_pessoas = pd.read_csv("Base_Pessoas.csv", sep=";")
    except:
        # Se o arquivo não existir, cria o DataFrame com o cabeçalho que definimos
        dados_pessoas = pd.DataFrame(columns=[
            "id_documento", "tipo_pessoa", "nome_razao", "nome_fantasia", 
            "rg_ie", "email", "telefone", "cep", "endereco", "numero", 
            "complemento", "bairro", "cidade", "uf", "categoria", 
            "limite_credito", "status", "data_cadastro"
        ])

    with st.form("form_pessoas"):
        st.subheader("1. Identificação Principal")
        col_id_1, col_id_2, col_id_3 = st.columns([2, 2, 2])
        
        with col_id_1:
            tipo_pessoa = st.selectbox("Tipo de Pessoa", ["Física", "Jurídica"])
            # Muda o rótulo do campo conforme a escolha
            label_doc = "CPF (Obrigatório)" if tipo_pessoa == "Física" else "CNPJ (Obrigatório)"
            id_documento = st.text_input(label_doc)
            
        with col_id_2:
            categoria = st.selectbox("Categoria", ["Cliente", "Fornecedor", "Transportadora", "Ambos"])
            status = st.selectbox("Status Inicial", ["Ativo", "Inativo", "Bloqueado"])
            
        with col_id_3:
            limite_credito = st.number_input("Limite de Crédito (R$)", min_value=0.0, step=100.0)

        st.divider()

        st.subheader("2. Dados Pessoais / Empresariais")
        col_dados_1, col_dados_2 = st.columns(2)
        
        with col_dados_1:
            label_nome = "Nome Completo" if tipo_pessoa == "Física" else "Razão Social"
            nome_razao = st.text_input(label_nome)
            nome_fantasia = st.text_input("Nome Fantasia (Se houver)")
            
        with col_dados_2:
            label_rg = "RG" if tipo_pessoa == "Física" else "Inscrição Estadual"
            rg_ie = st.text_input(label_rg)
            email = st.text_input("E-mail para contato/NFe")
            telefone = st.text_input("WhatsApp / Telefone")

        st.divider()

        st.subheader("3. Endereço")
        col_end_1, col_end_2, col_end_3 = st.columns([1, 2, 1])
        with col_end_1:
            cep = st.text_input("CEP")
            numero = st.text_input("Número")
        with col_end_2:
            endereco = st.text_input("Logradouro (Rua/Av)")
            complemento = st.text_input("Complemento")
        with col_end_3:
            bairro = st.text_input("Bairro")
            cidade = st.text_input("Cidade")
            uf = st.selectbox("UF", ["AC", "AL", "AP", "AM", "BA", "CE", "DF", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR", "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO"])

        botao_salvar_pessoa = st.form_submit_button("Finalizar Cadastro")

    # --- LÓGICA DE SALVAR PESSOA ---
    if botao_salvar_pessoa:
        erros_pessoa = []
        
        # 1. Validação de Documento e Nome
        if not id_documento: erros_pessoa.append("O campo CPF/CNPJ é obrigatório.")
        if not nome_razao: erros_pessoa.append(f"O campo {label_nome} é obrigatório.")
        
        # 2. Trava de Duplicidade
        if str(id_documento) in dados_pessoas["id_documento"].astype(str).tolist():
            erros_pessoa.append(f"Este documento ({id_documento}) já está cadastrado no sistema!")

        if len(erros_pessoa) > 0:
            for erro in erros_pessoa:
                st.error(erro)
        else:
            # 3. Criar nova linha
            nova_pessoa = pd.DataFrame({
                "id_documento": [id_documento],
                "tipo_pessoa": [tipo_pessoa],
                "nome_razao": [nome_razao],
                "nome_fantasia": [nome_fantasia],
                "rg_ie": [rg_ie],
                "email": [email],
                "telefone": [telefone],
                "cep": [cep],
                "endereco": [endereco],
                "numero": [numero],
                "complemento": [complemento],
                "bairro": [bairro],
                "cidade": [cidade],
                "uf": [uf],
                "categoria": [categoria],
                "limite_credito": [limite_credito],
                "status": [status],
                "data_cadastro": [datetime.now().strftime("%d/%m/%Y")]
            })

            # 4. Salvar no CSV
            dados_pessoas = pd.concat([dados_pessoas, nova_pessoa], ignore_index=True)
            dados_pessoas.to_csv("Base_Pessoas.csv", sep=";", index=False)
            
            st.success(f"✅ {tipo_pessoa} '{nome_razao}' cadastrada com sucesso!")
# --- 6. TELA DE CONSULTA DE PESSOAS ---
elif pagina == "Consultar Pessoa":
    st.title("Consulta de Clientes / Fornecedores")

    # 1. Carregar os dados
    try:
        dados_pessoas = pd.read_csv("Base_Pessoas.csv", sep=";")
    except:
        st.warning("Nenhuma base de pessoas encontrada. Cadastre alguém primeiro!")
        st.stop() # Para a execução aqui caso não haja arquivo

    # 2. Layout de Colunas
    col_filtros, col_resultado = st.columns([1, 3])

    with col_filtros:
        st.subheader("Filtros")
        filtro_doc = st.text_input("Buscar por CPF/CNPJ")
        filtro_nome = st.text_input("Buscar por Nome/Razão")
        
        filtro_cat = st.multiselect(
            "Filtrar Categoria", 
            ["Cliente", "Fornecedor", "Transportadora"],
            default=[]
        )
        
        st.divider()
        st.caption("A busca por nome funciona com termos parciais.")

    with col_resultado:
        # Criamos a cópia para filtrar
        df_p_filtrado = dados_pessoas.copy()

        # Lógica de Filtro em Tempo Real
        if filtro_doc:
            df_p_filtrado = df_p_filtrado[df_p_filtrado["id_documento"].astype(str).str.contains(filtro_doc, na=False)]
        
        if filtro_nome:
            df_p_filtrado = df_p_filtrado[df_p_filtrado["nome_razao"].astype(str).str.contains(filtro_nome, case=False, na=False)]
        
        if filtro_cat:
            # Filtra se a categoria está na lista selecionada no multiselect
            df_p_filtrado = df_p_filtrado[df_p_filtrado["categoria"].isin(filtro_cat)]

        # Seleção de Colunas para a Tabela (O que o usuário precisa ver rápido)
        colunas_ver = [
            "id_documento",
            "nome_razao",
            "categoria",
            "email",
            "telefone",
            "cidade",
            "status"
        ]

        try:
            exibicao_p = df_p_filtrado[colunas_ver]
            
            # Renomeando para ficar apresentável
            exibicao_p.columns = ["DOCUMENTO", "NOME / RAZÃO SOCIAL", "CATEGORIA", "E-MAIL", "CONTATO", "CIDADE", "STATUS"]

            st.subheader(f"Registros Encontrados ({len(exibicao_p)})")
            
            if len(exibicao_p) > 0:
                st.dataframe(
                    exibicao_p, 
                    use_container_width=True, 
                    hide_index=True
                )
                
                # Widget extra: Ver detalhes completos
                if len(exibicao_p) == 1:
                    st.info("💡 Apenas um registro encontrado. Você pode ver todos os dados dele na tabela acima arrastando a barra de rolagem.")
            else:
                st.info("Nenhuma pessoa encontrada com esses critérios.")
                
        except KeyError as e:
            st.error(f"Erro nas colunas do arquivo: {e}")

# --- 7. TELA DE PEDIDOS (VERSÃO FINAL COM DESCONTO/ACRÉSCIMO E FRETE) ---
elif pagina == "Criar Pedido":
    st.title("Central de Pedidos")

    # --- INICIALIZAÇÃO DE ESTADOS ---
    if "carrinho" not in st.session_state:
        st.session_state.carrinho = []
    if "cliente_selecionado" not in st.session_state:
        st.session_state.cliente_selecionado = None
    if "produto_selecionado" not in st.session_state:
        st.session_state.produto_selecionado = None

    # Lógica de ID Sequencial
    if os.path.exists("Base_Pedido.csv"):
        try:
            base_pedidos_temp = pd.read_csv("Base_Pedido.csv", sep=";")
            proximo_id = base_pedidos_temp["id_pedido"].max() + 1 if not base_pedidos_temp.empty else 1
        except:
            proximo_id = 1
    else:
        proximo_id = 1

    st.subheader(f"Pedido Nº: {proximo_id}")

    # --- FUNÇÕES DE BUSCA (DIALOGS) ---
    @st.dialog("Buscar Cliente")
    def buscar_cliente_pop():
        st.write("Pesquise e selecione o cliente.")
        filtro = st.text_input("Nome ou CPF/CNPJ")
        if os.path.exists("Base_Pessoas.csv"):
            df_p = pd.read_csv("Base_Pessoas.csv", sep=";")
            if filtro:
                # Filtro flexível para Nome ou Documento
                df_p = df_p[df_p["nome_razao"].str.contains(filtro, case=False) | 
                            df_p["id_documento"].astype(str).str.contains(filtro)]
            
            for _, row in df_p.head(10).iterrows():
                col1, col2 = st.columns([3, 1])
                col1.write(f"**{row['nome_razao']}** ({row['id_documento']})")
                if col2.button("Selecionar", key=f"sel_p_{row['id_documento']}"):
                    st.session_state.cliente_selecionado = row.to_dict()
                    st.rerun()
        else:
            st.error("Base de Pessoas não encontrada!")

    @st.dialog("Buscar Produto")
    def buscar_produto_pop():
        st.write("Pesquise o SKU ou Descrição")
        filtro = st.text_input("Palavra-chave")
        if os.path.exists("Base de Dados.csv"):
            df_prod = pd.read_csv("Base de Dados.csv", sep=";")
            if filtro:
                df_prod = df_prod[df_prod["descricao"].str.contains(filtro, case=False) | 
                                  df_prod["id_sku"].astype(str).str.contains(filtro)]
            
            for _, row in df_prod.head(10).iterrows():
                col1, col2 = st.columns([3, 1])
                col1.write(f"**{row['id_sku']}** - {row['descricao']}")
                if col2.button("Selecionar", key=f"sel_prod_{row['id_sku']}"):
                    st.session_state.produto_selecionado = row.to_dict()
                    st.rerun()
        else:
            st.error("Base de Produtos não encontrada!")

    # --- ÁREA 1: IDENTIFICAÇÃO DO CLIENTE ---
    with st.container(border=True):
        col_cli_1, col_cli_2 = st.columns([3, 1])
        with col_cli_1:
            doc_exibicao = st.session_state.cliente_selecionado['id_documento'] if st.session_state.cliente_selecionado else ""
            st.text_input("Cliente (Selecione na busca)", value=doc_exibicao, disabled=True)
        with col_cli_2:
            st.write("##")
            if st.button("Buscar Cliente", use_container_width=True):
                buscar_cliente_pop()

        if st.session_state.cliente_selecionado:
            c = st.session_state.cliente_selecionado
            st.success(f"**{c['nome_razao']}** | {c['cidade']}-{c['uf']} | Limite: R$ {c['limite_credito']}")

    # --- ÁREA 2: INCLUSÃO DE PRODUTOS ---
    with st.container(border=True):
        st.write("###Adicionar Itens")
        col_prod_1, col_prod_2, col_prod_3 = st.columns([2, 1, 1])
        
        with col_prod_1:
            sku_exibicao = st.session_state.produto_selecionado['id_sku'] if st.session_state.produto_selecionado else ""
            st.text_input("SKU (Selecione na busca)", value=sku_exibicao, disabled=True)
        with col_prod_2:
            st.write("##")
            if st.button("Buscar Produto", use_container_width=True):
                buscar_produto_pop()
        with col_prod_3:
            qtd = st.number_input("Quantidade", min_value=1, value=1)

        if st.session_state.produto_selecionado:
            p = st.session_state.produto_selecionado
            st.info(f"**Produto:** {p['descricao']} | **Preço Base:** R$ {p['valor_liquido']:.2f}")
            
            col_desc_1, col_desc_2 = st.columns(2)
            
            # AJUSTE: Removido min_value para permitir acréscimos (desconto negativo)
            desconto = col_desc_1.number_input(
                "Desconto (R$) - Use negativo para Acréscimo", 
                value=0.0, 
                step=1.0
            )
            
            valor_final_item = float(p['valor_liquido']) - desconto
            
            # Lógica visual para o Delta do Metric
            label_delta = "Acréscimo" if desconto < 0 else "Desconto"
            col_desc_2.metric("Preço Unitário Final", f"R$ {valor_final_item:.2f}", delta=f"{-desconto:.2f} ({label_delta})")

            if st.button("Adicionar ao Carrinho", use_container_width=True):
                # Evita duplicados no carrinho
                if any(item['sku'] == p['id_sku'] for item in st.session_state.carrinho):
                    st.warning("Este produto já está no carrinho!")
                else:
                    st.session_state.carrinho.append({
                        "sku": p['id_sku'],
                        "descricao": p['descricao'],
                        "qtd": qtd,
                        "valor_unit": valor_final_item,
                        "subtotal": valor_final_item * qtd
                    })
                    st.session_state.produto_selecionado = None
                    st.rerun()

    # --- ÁREA 3: REVISÃO E FINALIZAÇÃO ---
    if st.session_state.carrinho:
        st.divider()
        st.subheader("Resumo do Pedido")
        df_carrinho = pd.DataFrame(st.session_state.carrinho)
        st.dataframe(df_carrinho, use_container_width=True, hide_index=True)
        
        if st.button("Esvaziar Carrinho"):
            st.session_state.carrinho = []
            st.rerun()

        subtotal_itens = df_carrinho["subtotal"].sum()

        with st.form("finalizar_venda"):
            f1, f2, f3 = st.columns([2, 1, 1])
            
            with f1:
                tipo = st.selectbox("Tipo", ["ORÇAMENTO", "PEDIDO", "COTAÇÃO"])
                obs = st.text_area("Observações")
            with f2:
                frete = st.number_input("Frete (R$)", min_value=0.0, step=5.0)
                st.write(f"**Subtotal:** R$ {subtotal_itens:.2f}")
            with f3:
                total_final = subtotal_itens + frete
                st.metric("TOTAL GERAL", f"R$ {total_final:.2f}")

            if st.form_submit_button("CONFIRMAR E SALVAR", use_container_width=True):
                if not st.session_state.cliente_selecionado:
                    st.error("Selecione um cliente!")
                else:
                    novas_linhas = []
                    for item in st.session_state.carrinho:
                        novas_linhas.append({
                            "id_pedido": proximo_id,
                            "data_pedido": datetime.now().strftime("%d/%m/%Y %H:%M"),
                            "doc_cliente": st.session_state.cliente_selecionado['id_documento'],
                            "nome_cliente": st.session_state.cliente_selecionado['nome_razao'],
                            "sku_item": item['sku'],
                            "qtd": item['qtd'],
                            "valor_final": item['valor_unit'],
                            "frete_total": frete,
                            "tipo": tipo,
                            "observacao": obs
                        })
                    
                    df_save = pd.DataFrame(novas_linhas)
                    header_status = not os.path.exists("Base_Pedido.csv")
                    df_save.to_csv("Base_Pedido.csv", mode='a', sep=";", index=False, header=header_status)
                    
                    st.success("Pedido Gravado!")
                    st.session_state.carrinho = []
                    st.session_state.cliente_selecionado = None
                    st.rerun()

# --- 8. TELA DE CONSULTA DE PEDIDOS (VERSÃO COMPLETA COM FRETE) ---
elif pagina == "Consultar Pedido":
    st.title("Gestão e Consulta de Pedidos")

    # 1. Carregamento e Padronização de Dados
    try:
        df_pedidos = pd.read_csv("Base_Pedido.csv", sep=";")
        df_produtos = pd.read_csv("Base de Dados.csv", sep=";")
        df_pessoas = pd.read_csv("Base_Pessoas.csv", sep=";")
        
        # PADRONIZAÇÃO DE TIPOS (Evita o erro de Merge int64 vs Object)
        df_pedidos['sku_item'] = df_pedidos['sku_item'].astype(str)
        df_produtos['id_sku'] = df_produtos['id_sku'].astype(str)
        df_pedidos['doc_cliente'] = df_pedidos['doc_cliente'].astype(str)
        df_pessoas['id_documento'] = df_pessoas['id_documento'].astype(str)

        # Conversão de data para filtro
        df_pedidos['data_pedido_dt'] = pd.to_datetime(df_pedidos['data_pedido'], format="%d/%m/%Y %H:%M")
    except Exception as e:
        st.error(f"Erro ao carregar bases: {e}. Verifique se os arquivos existem.")
        st.stop()

    # 2. Layout Lateral (Filtros)
    col_filtros, col_detalhe = st.columns([1, 2.5])

    with col_filtros:
        st.subheader("Filtros")
        
        f_id_check = st.checkbox("Nº do Pedido")
        f_id = st.number_input("ID", min_value=1, step=1, disabled=not f_id_check)

        f_data_check = st.checkbox("Período")
        f_data_ini = st.date_input("Início", disabled=not f_data_check)
        f_data_fim = st.date_input("Fim", disabled=not f_data_check)

        f_cliente_check = st.checkbox("CPF/CNPJ")
        f_cliente = st.text_input("Documento", disabled=not f_cliente_check)

        f_sku_check = st.checkbox("SKU")
        f_sku = st.text_input("SKU do Item", disabled=not f_sku_check)

        # Lógica de Filtragem
        df_f = df_pedidos.copy()
        if f_id_check: df_f = df_f[df_f["id_pedido"] == f_id]
        if f_data_check:
            df_f = df_f[(df_f['data_pedido_dt'].dt.date >= f_data_ini) & 
                        (df_f['data_pedido_dt'].dt.date <= f_data_fim)]
        if f_cliente_check:
            df_f = df_f[df_f["doc_cliente"].str.contains(f_cliente)]
        if f_sku_check:
            df_f = df_f[df_f["sku_item"].str.contains(f_sku)]

        st.divider()
        lista_ids = df_f["id_pedido"].unique()
        if len(lista_ids) > 0:
            id_selecionado = st.selectbox("Selecione o Pedido", sorted(lista_ids, reverse=True))
        else:
            st.warning("Nenhum pedido encontrado.")
            id_selecionado = None

    # 3. Coluna Direita (Detalhes do Pedido)
    with col_detalhe:
        if id_selecionado:
            # Filtra itens e cruza com dados de produtos
            itens_pedido = df_pedidos[df_pedidos["id_pedido"] == id_selecionado]
            itens_completos = itens_pedido.merge(
                df_produtos[['id_sku', 'fornecedor', 'descricao', 'preco_custo']], 
                left_on='sku_item', right_on='id_sku', how='left'
            )

            # Busca dados do cliente
            doc_do_pedido = str(itens_pedido.iloc[0]["doc_cliente"])
            # Filtro seguro para encontrar o cliente
            dados_cli_df = df_pessoas[df_pessoas["id_documento"] == doc_do_pedido]
            
            if not dados_cli_df.empty:
                cli = dados_cli_df.iloc[0]
                endereco_completo = f"{cli['endereco']}, {cli['numero']} - {cli['cidade']}/{cli['uf']}"
                cat_cli = cli['categoria']
            else:
                endereco_completo = "Cliente não encontrado na base de Pessoas"
                cat_cli = "N/A"

            # --- CÁLCULOS FINANCEIROS ---
            valor_frete = float(itens_pedido.iloc[0]["frete_total"])
            subtotal_itens = (itens_completos["valor_final"] * itens_completos["qtd"]).sum()
            custo_total_itens = (itens_completos["preco_custo"] * itens_completos["qtd"]).sum()
            
            total_geral = subtotal_itens + valor_frete
            lucro_total = subtotal_itens - custo_total_itens # Frete geralmente não entra no lucro bruto

            # --- CABEÇALHO SUPERIOR ---
            # --- CABEÇALHO SUPERIOR REORGANIZADO ---
            with st.container(border=True):
                # Informações do Cliente e Pedido
                c1, c2 = st.columns([2, 1])
                with c1:
                    st.markdown(f"###{itens_pedido.iloc[0]['nome_cliente']}")
                    st.caption(f"{endereco_completo} | {cat_cli}")
                with c2:
                    st.markdown(f"**Data:** {itens_pedido.iloc[0]['data_pedido']}")
                    st.markdown(f"**Tipo:** `{itens_pedido.iloc[0]['tipo']}`")
                
                st.divider()
                
                # Métricas Financeiras em duas linhas para não cortar os números
                col_m1, col_m2 = st.columns(2)
                with col_m1:
                    st.metric("Custo Total Itens", f"R$ {custo_total_itens:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
                    st.metric("Frete do Pedido", f"R$ {valor_frete:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
                
                with col_m2:
                    st.metric("Lucro Bruto Estimado", f"R$ {lucro_total:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."), delta_color="normal")
                    # O Valor Total ganha um destaque maior aqui
                    st.subheader("VALOR TOTAL")
                    st.title(f"R$ {total_geral:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))

            # --- TABELA INFERIOR ---
            st.subheader("Itens do Pedido")
            tabela_show = pd.DataFrame({
                "FORNECEDOR": itens_completos["fornecedor"],
                "SKU": itens_completos["sku_item"],
                "DESCRIÇÃO": itens_completos["descricao"],
                "QTD": itens_completos["qtd"],
                "CUSTO UN.": itens_completos["preco_custo"],
                "VALOR UN.": itens_completos["valor_final"],
                "TOTAL": itens_completos["valor_final"] * itens_completos["qtd"]
            })
            st.dataframe(tabela_show, use_container_width=True, hide_index=True)
            
            if str(itens_pedido.iloc[0]['observacao']) != 'nan':
                st.info(f"**Observações:** {itens_pedido.iloc[0]['observacao']}")
        else:
            st.info("Selecione um pedido à esquerda para ver os detalhes.")
# --- 9. TELA DE FORMALIZAÇÃO DE PROPOSTA (CORRIGIDA) ---
elif pagina == "Formalizacao":
    st.title("Formalização de Proposta")
    
    # 1. Carregamento e Padronização
    try:
        df_pedidos = pd.read_csv("Base_Pedido.csv", sep=";")
        df_produtos = pd.read_csv("Base de Dados.csv", sep=";")
        df_pessoas = pd.read_csv("Base_Pessoas.csv", sep=";")
        
        # Conversão de tipos para garantir o cruzamento (merge)
        df_pedidos['id_pedido'] = df_pedidos['id_pedido'].astype(int)
        df_pedidos['sku_item'] = df_pedidos['sku_item'].astype(str)
        df_produtos['id_sku'] = df_produtos['id_sku'].astype(str)
        df_pessoas['id_documento'] = df_pessoas['id_documento'].astype(str)
        
        lista_pedidos = sorted(df_pedidos["id_pedido"].unique(), reverse=True)
    except Exception as e:
        st.error(f"Erro ao carregar bases: {e}")
        st.stop()

    # 2. Seleção do Pedido
    id_escolhido = st.selectbox("Selecione o Número do Pedido", lista_pedidos, index=None, placeholder="Escolha um pedido...")

    if id_escolhido:
        # Filtra dados do Pedido
        dados_venda = df_pedidos[df_pedidos["id_pedido"] == id_escolhido]
        doc_cliente = str(dados_venda.iloc[0]["doc_cliente"])
        
        # AJUSTE NAS COLUNAS: Usando apenas o que existe no seu CSV (id_sku, descricao, marca, preco_custo)
        itens_completos = dados_venda.merge(
            df_produtos[['id_sku', 'descricao', 'marca', 'preco_custo']], 
            left_on='sku_item', right_on='id_sku', how='left'
        )
        
        # Dados do Cliente
        cliente_info = df_pessoas[df_pessoas["id_documento"] == doc_cliente].iloc[0]

        # --- PAINEL DE CONFERÊNCIA ---
        with st.container(border=True):
            st.subheader(f"Resumo: Pedido #{id_escolhido}")
            c1, c2 = st.columns(2)
            with c1:
                st.write(f"**Razão Social:** {cliente_info['nome_razao']}")
                st.write(f"**Data:** {dados_venda.iloc[0]['data_pedido']}")
                st.caption(f"📍 {cliente_info['endereco']}, {cliente_info['numero']} - {cliente_info['cidade']}/{cliente_info['uf']}")
            with c2:
                custo_total = (itens_completos['preco_custo'] * itens_completos['qtd']).sum()
                # Soma subtotal dos itens + o frete único do pedido
                venda_total = (itens_completos['valor_final'] * itens_completos['qtd']).sum() + float(dados_venda.iloc[0]['frete_total'])
                st.metric("Total Venda (c/ Frete)", f"R$ {venda_total:.2f}")
                st.write(f"**Custo Total Est.:** R$ {custo_total:.2f}")

            # Exibição da Observação do Pedido
            obs_pedido = dados_venda.iloc[0]['observacao']
            if pd.notna(obs_pedido) and str(obs_pedido).lower() != 'nan':
                st.warning(f"📝 **Observação do Pedido:** {obs_pedido}")

            st.write("**Itens Selecionados:**")
            st.dataframe(itens_completos[['sku_item', 'descricao', 'marca', 'qtd', 'valor_final']], use_container_width=True, hide_index=True)

        st.divider()

        # 3. Inputs Manuais para o Documento
        st.subheader("Dados Adicionais para Proposta_Modelo")
        with st.form("form_formalizacao"):
            f1, f2 = st.columns(2)
            with f1:
                n_pregao = st.text_input("Nº do Pregão / Processo")
                validade = st.text_input("Validade da Proposta (ex: 60 dias)")
            with f2:
                prazo = st.text_input("Prazo de Entrega (ex: 15 dias úteis)")
                contato_doc = st.text_input("Pessoa de Contato", value=cliente_info['nome_razao'])
            
            especificacoes = st.text_area("Especificações Técnicas Solicitadas")
            
            botao_gerar = st.form_submit_button("Gerar Proposta (Word)", use_container_width=True)

        # 4. Geração do Word
        if botao_gerar:
            if not all([n_pregao, validade, prazo, especificacoes]):
                st.error("Preencha todos os campos obrigatórios para gerar o documento.")
            else:
                try:
                    from num2words import num2words
                    doc = Document("Proposta_Modelo.docx")
                    
                    # Valor por extenso
                    valor_extenso = num2words(venda_total, lang='pt_BR', to='currency').upper()

                    # Substituição de Tags
                    subs = {
                        "[Razao_UASG]": cliente_info['nome_razao'],
                        "[N_pregao]": n_pregao,
                        "[Esp_solicitadas]": especificacoes,
                        "[Validade_Proposta]": validade,
                        "[Prazo_entrega]": prazo,
                        "[Endereco_Cliente]": f"{cliente_info['endereco']}, {cliente_info['numero']}",
                        "[Contato_Cliente]": contato_doc,
                        "MIL QUINHENTOS E QUARENTA REAIS": valor_extenso
                    }

                    for p in doc.paragraphs:
                        for tag, val in subs.items():
                            if tag in p.text:
                                p.text = p.text.replace(tag, str(val))

                    # Preenchimento da Tabela
                    if doc.tables:
                        tabela = doc.tables[0]
                        for i, it in itens_completos.iterrows():
                            cells = tabela.add_row().cells
                            cells[0].text = str(i + 1)
                            cells[1].text = str(it['descricao'])
                            cells[2].text = str(it['marca'])
                            cells[3].text = "---" # PartNumber não existe no seu CSV, deixamos fixo ou vazio
                            cells[4].text = str(it['qtd'])
                            cells[5].text = f"R$ {it['valor_final']:.2f}"
                            cells[6].text = f"R$ {(it['valor_final'] * it['qtd']):.2f}"

                    nome_final = f"Proposta_{id_escolhido}.docx"
                    doc.save(nome_final)
                    
                    with open(nome_final, "rb") as f:
                        st.download_button("Baixar Proposta Gerada", f, file_name=nome_final)
                    st.success("Documento gerado!")

                except Exception as e:

                    st.error(f"Erro na geração do documento: {e}")

